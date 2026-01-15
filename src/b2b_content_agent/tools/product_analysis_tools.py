"""Product Analyst Tools for extracting and analyzing product information."""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Any
from urllib.parse import urlparse

from crewai.tools import BaseTool
from pydantic import BaseModel, Field
import PyPDF2
import docx
import requests
from bs4 import BeautifulSoup


class ProductInfo(BaseModel):
    """Structured product information extracted from various sources."""
    
    product_name: Optional[str] = Field(None, description="Name of the product or service")
    description: Optional[str] = Field(None, description="Product description")
    features: List[str] = Field(default_factory=list, description="List of product features")
    benefits: List[str] = Field(default_factory=list, description="List of product benefits")
    use_cases: List[str] = Field(default_factory=list, description="List of use cases")
    target_market: List[str] = Field(default_factory=list, description="Target market segments")
    pricing_info: Optional[str] = Field(None, description="Pricing information if available")
    technical_specs: Dict[str, Any] = Field(default_factory=dict, description="Technical specifications")
    industry: Optional[str] = Field(None, description="Industry classification")
    competitors: List[str] = Field(default_factory=list, description="Known competitors")
    unique_value_proposition: Optional[str] = Field(None, description="What makes it unique")
    raw_content: str = Field("", description="Raw extracted content")


class DocumentParserTool(BaseTool):
    """Tool for parsing various document formats (PDF, DOCX, TXT, MD)."""
    
    name: str = "Document Parser"
    description: str = (
        "Parses product documentation from various file formats including "
        "PDF, DOCX, TXT, and Markdown files. Extracts text content and "
        "basic structure for further analysis."
    )
    
    def _run(self, file_path: str) -> str:
        """
        Parse a document and extract its content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content from the document
        """
        path = Path(file_path)
        
        if not path.exists():
            return f"Error: File not found at {file_path}"
        
        try:
            file_extension = path.suffix.lower()
            
            if file_extension == '.pdf':
                return self._parse_pdf(path)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_docx(path)
            elif file_extension in ['.txt', '.md', '.markdown']:
                return self._parse_text(path)
            else:
                return f"Error: Unsupported file format {file_extension}"
                
        except Exception as e:
            return f"Error parsing document: {str(e)}"
    
    def _parse_pdf(self, path: Path) -> str:
        """Extract text from PDF file."""
        text_content = []
        
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
        
        return "\n\n".join(text_content)
    
    def _parse_docx(self, path: Path) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(path)
        
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text_content.append(" | ".join(row_text))
        
        return "\n\n".join(text_content)
    
    def _parse_text(self, path: Path) -> str:
        """Extract text from plain text or markdown files."""
        with open(path, 'r', encoding='utf-8') as file:
            return file.read()


class WebScraperTool(BaseTool):
    """Tool for scraping product information from websites."""
    
    name: str = "Web Scraper"
    description: str = (
        "Scrapes product information from websites including product pages, "
        "landing pages, and documentation sites. Extracts relevant content "
        "while filtering out navigation and boilerplate."
    )
    
    def _run(self, url: str) -> str:
        """
        Scrape content from a URL.
        
        Args:
            url: Website URL to scrape
            
        Returns:
            Extracted and cleaned content from the website
        """
        try:
            # Validate URL
            parsed_url = urlparse(url)
            if not parsed_url.scheme or not parsed_url.netloc:
                return f"Error: Invalid URL format: {url}"
            
            # Set headers to mimic browser request
            headers = {
                'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5',
            }
            
            # Fetch the page
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script, style, and nav elements
            for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
                tag.decompose()
            
            # Extract title
            title = soup.find('title')
            title_text = title.get_text(strip=True) if title else ""
            
            # Extract meta description
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            description = meta_desc.get('content', '') if meta_desc else ""
            
            # Extract main content
            # Try to find main content area
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile('content|main'))
            
            if main_content:
                text_content = main_content.get_text(separator='\n', strip=True)
            else:
                # Fallback to body
                text_content = soup.body.get_text(separator='\n', strip=True) if soup.body else ""
            
            # Clean up excessive whitespace
            text_content = re.sub(r'\n\s*\n', '\n\n', text_content)
            
            result = f"URL: {url}\n"
            if title_text:
                result += f"Title: {title_text}\n"
            if description:
                result += f"Description: {description}\n"
            result += f"\n--- Content ---\n{text_content}"
            
            return result
            
        except requests.Timeout:
            return f"Error: Request timeout while accessing {url}"
        except requests.RequestException as e:
            return f"Error: Failed to fetch {url}: {str(e)}"
        except Exception as e:
            return f"Error: Failed to parse content from {url}: {str(e)}"


class ProductAnalyzerTool(BaseTool):
    """Tool for analyzing and structuring product information from raw content."""
    
    name: str = "Product Analyzer"
    description: str = (
        "Analyzes raw product content and extracts structured information "
        "including features, benefits, use cases, target market, and more. "
        "Uses pattern matching and AI to identify key product attributes."
    )
    
    def _run(self, content: str) -> str:
        """
        Analyze raw content and extract structured product information.
        
        Args:
            content: Raw text content about a product
            
        Returns:
            Structured analysis of the product information
        """
        product_info = ProductInfo(raw_content=content)
        
        # Extract product name (usually in first few lines or title)
        lines = content.split('\n')
        for line in lines[:10]:
            if line.strip() and len(line.strip()) < 100:
                # Likely a title/heading
                if not product_info.product_name:
                    product_info.product_name = line.strip()
                    break
        
        # Pattern matching for common sections
        self._extract_features(content, product_info)
        self._extract_benefits(content, product_info)
        self._extract_use_cases(content, product_info)
        self._extract_target_market(content, product_info)
        self._extract_pricing(content, product_info)
        self._extract_technical_specs(content, product_info)
        
        # Format output
        result = self._format_product_info(product_info)
        return result
    
    def _extract_features(self, content: str, product_info: ProductInfo):
        """Extract product features from content."""
        patterns = [
            r'(?:feature|capability|functionality)s?:?\s*\n((?:[-•*]\s*.+\n?)+)',
            r'(?:what it does|key features|main features):?\s*\n((?:[-•*]\s*.+\n?)+)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                items = re.findall(r'[-•*]\s*(.+)', match.group(1))
                product_info.features.extend([item.strip() for item in items if item.strip()])
    
    def _extract_benefits(self, content: str, product_info: ProductInfo):
        """Extract product benefits from content."""
        patterns = [
            r'(?:benefit|advantage|value)s?:?\s*\n((?:[-•*]\s*.+\n?)+)',
            r'(?:why choose|why use|advantages):?\s*\n((?:[-•*]\s*.+\n?)+)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                items = re.findall(r'[-•*]\s*(.+)', match.group(1))
                product_info.benefits.extend([item.strip() for item in items if item.strip()])
    
    def _extract_use_cases(self, content: str, product_info: ProductInfo):
        """Extract use cases from content."""
        patterns = [
            r'(?:use case|application|scenario)s?:?\s*\n((?:[-•*]\s*.+\n?)+)',
            r'(?:who can use|ideal for|perfect for):?\s*\n((?:[-•*]\s*.+\n?)+)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                items = re.findall(r'[-•*]\s*(.+)', match.group(1))
                product_info.use_cases.extend([item.strip() for item in items if item.strip()])
    
    def _extract_target_market(self, content: str, product_info: ProductInfo):
        """Extract target market information."""
        # Look for common target market keywords
        target_keywords = [
            'enterprise', 'small business', 'startup', 'b2b', 'b2c',
            'developers', 'marketers', 'sales teams', 'managers',
            'professionals', 'individuals', 'organizations'
        ]
        
        content_lower = content.lower()
        for keyword in target_keywords:
            if keyword in content_lower:
                product_info.target_market.append(keyword.title())
    
    def _extract_pricing(self, content: str, product_info: ProductInfo):
        """Extract pricing information."""
        # Look for pricing patterns
        pricing_patterns = [
            r'\$\d+(?:,\d{3})*(?:\.\d{2})?(?:/month|/mo|/year|/yr)?',
            r'(?:free|freemium|subscription|license)',
            r'(?:pricing|price|cost):?\s*(.+)',
        ]
        
        for pattern in pricing_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            if matches:
                product_info.pricing_info = ', '.join(str(m) for m in matches[:5])
                break
    
    def _extract_technical_specs(self, content: str, product_info: ProductInfo):
        """Extract technical specifications."""
        # Look for common tech spec patterns
        spec_patterns = {
            'platform': r'(?:platform|os|operating system):?\s*([^\n]+)',
            'languages': r'(?:languages?|programming languages?):?\s*([^\n]+)',
            'integration': r'(?:integrates? with|supports?):?\s*([^\n]+)',
            'api': r'(?:api|rest api|graphql):?\s*([^\n]+)',
        }
        
        for key, pattern in spec_patterns.items():
            matches = re.findall(pattern, content, re.IGNORECASE)
            if matches:
                product_info.technical_specs[key] = matches[0].strip()
    
    def _format_product_info(self, product_info: ProductInfo) -> str:
        """Format product information into a readable string."""
        sections = []
        
        if product_info.product_name:
            sections.append(f"=== PRODUCT: {product_info.product_name} ===\n")
        
        if product_info.features:
            sections.append("FEATURES:")
            for i, feature in enumerate(product_info.features, 1):
                sections.append(f"  {i}. {feature}")
            sections.append("")
        
        if product_info.benefits:
            sections.append("BENEFITS:")
            for i, benefit in enumerate(product_info.benefits, 1):
                sections.append(f"  {i}. {benefit}")
            sections.append("")
        
        if product_info.use_cases:
            sections.append("USE CASES:")
            for i, use_case in enumerate(product_info.use_cases, 1):
                sections.append(f"  {i}. {use_case}")
            sections.append("")
        
        if product_info.target_market:
            sections.append(f"TARGET MARKET: {', '.join(product_info.target_market)}\n")
        
        if product_info.pricing_info:
            sections.append(f"PRICING: {product_info.pricing_info}\n")
        
        if product_info.technical_specs:
            sections.append("TECHNICAL SPECS:")
            for key, value in product_info.technical_specs.items():
                sections.append(f"  - {key.title()}: {value}")
            sections.append("")
        
        return "\n".join(sections)


class CompetitorAnalyzerTool(BaseTool):
    """Tool for identifying and analyzing competitors from product information."""
    
    name: str = "Competitor Analyzer"
    description: str = (
        "Identifies potential competitors based on product information and "
        "industry classification. Can search for and analyze competitor products."
    )
    
    def _run(self, product_description: str, industry: str = None) -> str:
        """
        Analyze potential competitors based on product description.
        
        Args:
            product_description: Description of the product
            industry: Optional industry classification
            
        Returns:
            Analysis of potential competitors
        """
        # This is a placeholder for now - would integrate with search APIs in production
        result = [
            "=== COMPETITOR ANALYSIS ===\n",
            "Note: This is a basic analysis. For comprehensive competitor research,",
            "integrate with search APIs, market research databases, and industry reports.\n",
        ]
        
        # Extract key terms from product description
        keywords = self._extract_keywords(product_description)
        
        if keywords:
            result.append(f"Key Product Keywords: {', '.join(keywords)}\n")
            result.append("Recommended Competitor Research Strategy:")
            result.append(f"  1. Search for: '{' '.join(keywords[:3])}'")
            result.append(f"  2. Look for alternatives to similar products")
            result.append(f"  3. Research companies in {industry or 'related industries'}")
            result.append(f"  4. Check industry reports and analyst reviews\n")
        
        return "\n".join(result)
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extract key terms from product description."""
        # Simple keyword extraction - would use NLP in production
        common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
        
        words = re.findall(r'\b[a-z]{4,}\b', text.lower())
        keywords = [w for w in words if w not in common_words]
        
        # Return top 10 most frequent
        from collections import Counter
        word_counts = Counter(keywords)
        return [word for word, count in word_counts.most_common(10)]
